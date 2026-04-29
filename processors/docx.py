"""DOCX → Markdown：用 python-docx 提取。"""
from pathlib import Path
from docx import Document


def process_docx(docx_path: str, output_path: str, config: dict | None = None) -> dict:
    doc = Document(docx_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 根据样式添加标题标记
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                parts.append(f"# {text}")
            elif "heading 2" in style_name:
                parts.append(f"## {text}")
            elif "heading 3" in style_name:
                parts.append(f"### {text}")
            else:
                parts.append(text)
        else:
            parts.append("")

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            parts.append(" | ".join(cells))
        parts.append("")

    text = "\n".join(parts)
    Path(output_path).write_text(text, "utf-8", newline="\r\n")
    return {"status": "ok", "chars": len(text)}
